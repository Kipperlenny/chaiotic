"""Module for extracting structured content from documents."""

import os
import re
import zipfile
from lxml import etree
from typing import List, Dict, Any

from .document_reader import DOCX_AVAILABLE, ODF_AVAILABLE, read_document

def extract_structured_content(file_path):
    """Extract structured content from a document, preserving paragraph metadata.
    
    Args:
        file_path: Path to the document file.
        
    Returns:
        List of dictionaries containing structured content, where each dictionary has:
        - 'id': A unique identifier for the element
        - 'type': The element type (paragraph, heading, list, table)
        - 'content': The text content of the element
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.docx':
        return extract_structured_docx(file_path)
    elif file_ext == '.odt':
        return extract_structured_odt(file_path)
    elif file_ext == '.txt':
        return extract_structured_text(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")

def extract_structured_docx(file_path):
    """Extract structured content from a DOCX file."""
    try:
        if not DOCX_AVAILABLE:
            print("python-docx library not available. Using basic XML parsing for structured DOCX.")
            print("For better results, install: pip install python-docx")
            return extract_structured_docx_xml(file_path)
        
        from docx import Document
        doc = Document(file_path)
        structured_content = []
        
        # Process paragraphs
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue  # Skip empty paragraphs
                
            # Determine paragraph type based on style
            para_type = "paragraph"
            if para.style.name.startswith('Heading'):
                para_type = f"heading-{para.style.name.split(' ')[-1]}"
            elif para.style.name == 'List Paragraph':
                para_type = "list-item"
                
            structured_content.append({
                'id': f'p{i+1}',
                'type': para_type,
                'content': para.text
            })
        
        # Process tables
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    # Combine all paragraphs in a cell
                    cell_text = '\n'.join(p.text for p in cell.paragraphs if p.text.strip())
                    if cell_text.strip():
                        structured_content.append({
                            'id': f't{t_idx+1}r{r_idx+1}c{c_idx+1}',
                            'type': 'table-cell',
                            'content': cell_text
                        })
        
        return structured_content
    except Exception as e:
        print(f"Error extracting structured content from DOCX: {e}")
        # Fall back to XML parsing
        try:
            return extract_structured_docx_xml(file_path)
        except Exception as e2:
            print(f"XML fallback failed: {e2}")
            # Last resort - use simple text
            content, _, _ = read_document(file_path)
            return extract_structured_text_from_content(content)

def extract_structured_docx_xml(file_path):
    """Extract structured content from a DOCX file using XML parsing."""
    import zipfile
    import xml.etree.ElementTree as ET
    
    try:
        # Try to use lxml for better XML handling
        import lxml.etree as LET
        parser = LET
    except ImportError:
        parser = ET
    
    structured_content = []
    
    try:
        with zipfile.ZipFile(file_path) as docx:
            # Read document.xml
            with docx.open('word/document.xml') as content:
                tree = parser.parse(content)
                root = tree.getroot()
                
                # Define namespaces
                ns = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }
                
                # Extract paragraphs
                for i, para in enumerate(root.findall('.//w:p', ns)):
                    para_text = []
                    style = None
                    
                    # Get paragraph style
                    style_elem = para.find('.//w:pStyle', ns)
                    if style_elem is not None:
                        style = style_elem.get('{%s}val' % ns['w'])
                    
                    # Get text runs
                    for run in para.findall('.//w:t', ns):
                        if run.text:
                            para_text.append(run.text)
                    
                    text = ''.join(para_text).strip()
                    if text:
                        para_type = "paragraph"
                        if style and style.startswith('Heading'):
                            level = style.replace('Heading', '')
                            try:
                                level = int(level)
                                para_type = f"heading-{level}"
                            except ValueError:
                                para_type = "heading"
                        
                        structured_content.append({
                            'id': f'p{i+1}',
                            'type': para_type,
                            'content': text
                        })
                
                # Extract tables
                for t_idx, table in enumerate(root.findall('.//w:tbl', ns)):
                    for r_idx, row in enumerate(table.findall('.//w:tr', ns)):
                        for c_idx, cell in enumerate(row.findall('.//w:tc', ns)):
                            cell_text = []
                            for para in cell.findall('.//w:p', ns):
                                for run in para.findall('.//w:t', ns):
                                    if run.text:
                                        cell_text.append(run.text)
                            
                            text = ' '.join(cell_text).strip()
                            if text:
                                structured_content.append({
                                    'id': f't{t_idx+1}r{r_idx+1}c{c_idx+1}',
                                    'type': 'table-cell',
                                    'content': text
                                })
        
        return structured_content
        
    except Exception as e:
        print(f"Error extracting DOCX XML content: {e}")
        return []

def extract_structured_odt(file_path):
    """Extract structured content from an ODT file."""
    try:
        if not ODF_AVAILABLE:
            print("odfpy library not available. Using basic XML parsing for structured ODT.")
            return extract_structured_odt_xml(file_path)
        
        from odf.opendocument import load
        from odf.text import P, H, Span
        from odf.table import Table, TableRow, TableCell
        
        doc = load(file_path)
        structured_content = []
        
        def get_text_content(element):
            """Extract text from an ODT element."""
            parts = []
            for child in element.childNodes:
                if hasattr(child, 'data'):
                    parts.append(child.data)
                elif isinstance(child, Span):
                    parts.append(get_text_content(child))
            return ''.join(parts)
        
        # Process paragraphs and headings
        idx = 0
        for element in doc.text.childNodes:
            if isinstance(element, (P, H)):
                text = get_text_content(element)
                if text.strip():
                    idx += 1
                    structured_content.append({
                        'id': f'p{idx}',
                        'type': 'heading' if isinstance(element, H) else 'paragraph',
                        'content': text
                    })
            elif isinstance(element, Table):
                for r_idx, row in enumerate(element.getElementsByType(TableRow)):
                    for c_idx, cell in enumerate(row.getElementsByType(TableCell)):
                        text = get_text_content(cell)
                        if text.strip():
                            idx += 1
                            structured_content.append({
                                'id': f't{len(structured_content)+1}r{r_idx+1}c{c_idx+1}',
                                'type': 'table-cell',
                                'content': text
                            })
        
        return structured_content
        
    except Exception as e:
        print(f"Error extracting ODT content: {e}")
        return extract_structured_odt_xml(file_path)

def extract_structured_odt_xml(file_path):
    """Extract structured content from an ODT file using XML parsing."""
    import zipfile
    import xml.etree.ElementTree as ET
    
    try:
        # Try to use lxml for better XML handling
        import lxml.etree as LET
        parser = LET
    except ImportError:
        parser = ET
    
    structured_content = []
    
    try:
        with zipfile.ZipFile(file_path) as odt:
            # Read content.xml
            with odt.open('content.xml') as content:
                tree = parser.parse(content)
                root = tree.getroot()
                
                # Define namespaces
                ns = {
                    'office': 'urn:oasis:names:tc:opendocument:xmlns:office:1.0',
                    'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0',
                    'table': 'urn:oasis:names:tc:opendocument:xmlns:table:1.0'
                }
                
                # Process paragraphs and headings
                idx = 0
                for element in root.findall('.//text:p', ns) + root.findall('.//text:h', ns):
                    text = ''.join(element.itertext()).strip()
                    if text:
                        idx += 1
                        structured_content.append({
                            'id': f'p{idx}',
                            'type': 'heading' if element.tag.endswith('}h') else 'paragraph',
                            'content': text
                        })
                
                # Process tables
                for t_idx, table in enumerate(root.findall('.//table:table', ns)):
                    for r_idx, row in enumerate(table.findall('.//table:table-row', ns)):
                        for c_idx, cell in enumerate(row.findall('.//table:table-cell', ns)):
                            text = ''.join(cell.itertext()).strip()
                            if text:
                                structured_content.append({
                                    'id': f't{t_idx+1}r{r_idx+1}c{c_idx+1}',
                                    'type': 'table-cell',
                                    'content': text
                                })
        
        return structured_content
        
    except Exception as e:
        print(f"Error extracting ODT XML content: {e}")
        return []

def extract_structured_text(file_path):
    """Extract structured content from a plain text file."""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
    except UnicodeDecodeError:
        # Try with a different encoding if UTF-8 fails
        with open(file_path, 'r', encoding='latin-1') as f:
            content = f.read()
    
    return extract_structured_text_from_content(content)

def extract_structured_text_from_content(content):
    """Extract structured content from plain text content."""
    structured_content = []
    
    # Split into paragraphs
    paragraphs = re.split(r'\n\s*\n', content)
    
    for i, para in enumerate(paragraphs):
        para = para.strip()
        if not para:
            continue
        
        # Try to identify headings by looking for short single lines with capital letters
        lines = para.split('\n')
        if len(lines) == 1 and len(para) < 100 and para.isupper():
            structured_content.append({
                'id': f'p{i+1}',
                'type': 'heading-1',
                'content': para
            })
        elif len(lines) == 1 and len(para) < 100 and para[0].isupper() and not para.endswith('.'):
            structured_content.append({
                'id': f'p{i+1}',
                'type': 'heading-2',
                'content': para
            })
        else:
            # Regular paragraph
            structured_content.append({
                'id': f'p{i+1}',
                'type': 'paragraph',
                'content': para
            })
    
    return structured_content