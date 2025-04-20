"""Document handler module for reading and writing documents."""

import os
import json
from typing import List, Dict, Any, Optional, Tuple
from datetime import datetime

from .document_reader import read_document, read_docx, read_odt, read_text_file
from .document_extractor import extract_structured_content
from .docx_handler import (
    apply_corrections_to_docx,
    copy_docx_formatting,
    create_structured_docx
)
from .odt_handler import (
    apply_corrections_to_odt,
    create_structured_odt
)

def save_document(file_path: str, corrections: Dict[str, Any], 
                 original_doc: Any = None, is_docx: bool = True) -> Tuple[Optional[str], Optional[str], Optional[str]]:
    """Save document with corrections while preserving formatting.
    
    Args:
        file_path: Path to the original document
        corrections: Dictionary with corrections and metadata
        original_doc: Original document object if available
        is_docx: Whether the document is DOCX (True) or ODT (False)
    
    Returns:
        Tuple of (json_path, text_path, doc_path), any may be None on error
    """
    # Generate output filenames with timestamp
    file_dir = os.path.dirname(file_path) or '.'
    file_base = os.path.splitext(os.path.basename(file_path))[0]
    file_ext = os.path.splitext(file_path)[1]
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    
    json_path = os.path.join(file_dir, f"{file_base}_corrections_{timestamp}.json")
    text_path = os.path.join(file_dir, f"{file_base}_corrections_{timestamp}.txt")
    doc_path = os.path.join(file_dir, f"{file_base}_corrected_{timestamp}{file_ext}")
    
    try:
        # Normalize corrections to list format
        if isinstance(corrections, dict):
            if 'corrections' in corrections:
                corr_list = corrections['corrections']
                corrected_full_text = corrections.get('corrected_full_text', '')
            else:
                corr_list = [corrections] if corrections else []
                corrected_full_text = ''
        elif isinstance(corrections, list):
            corr_list = corrections
            corrected_full_text = ''
        else:
            corr_list = [corrections] if corrections else []
            corrected_full_text = ''
        
        # Save JSON file with metadata
        try:
            with open(json_path, 'w', encoding='utf-8') as f:
                json_data = {
                    'corrections': corr_list,
                    'metadata': {
                        'timestamp': datetime.now().isoformat(),
                        'original_file': file_path
                    }
                }
                if corrected_full_text:
                    json_data['corrected_full_text'] = corrected_full_text
                
                json.dump(json_data, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"Error saving JSON: {e}")
            json_path = None
        
        # Save text file with human-readable corrections
        try:
            with open(text_path, 'w', encoding='utf-8') as f:
                f.write("CORRECTIONS:\n\n")
                for i, corr in enumerate(corr_list, 1):
                    if isinstance(corr, dict):
                        f.write(f"{i}. Original: {corr.get('original', '')}\n")
                        f.write(f"   Corrected: {corr.get('corrected', '')}\n")
                        if 'explanation' in corr:
                            f.write(f"   Explanation: {corr['explanation']}\n")
                    elif isinstance(corr, str):
                        f.write(f"{i}. Correction: {corr}\n")
                    f.write("\n")
                
                # Add the full corrected text if available
                if corrected_full_text:
                    f.write("\n\nCORRECTED FULL TEXT:\n\n")
                    f.write(corrected_full_text)
        except Exception as e:
            print(f"Error saving text file: {e}")
            text_path = None
        
        # Apply corrections to document
        try:
            if is_docx:
                from .docx_handler import apply_corrections_to_docx
                doc_path = apply_corrections_to_docx(file_path, corr_list, doc_path)
            else:
                from .odt_handler import apply_corrections_to_odt
                try:
                    doc_path = apply_corrections_to_odt(file_path, corr_list, doc_path)
                except Exception as e:
                    print(f"Error saving ODT: {e}")
                    import traceback
                    traceback.print_exc()
                    text_path = create_fallback_document(corr_list, doc_path, is_docx)
                    doc_path = None
        except Exception as e:
            print(f"Error saving document: {e}")
            doc_path = None
    
    except Exception as e:
        print(f"Error in save_document: {e}")
        import traceback
        traceback.print_exc()
        return None, None, None
    
    return json_path, text_path, doc_path

def save_document_content(content: str, file_path: str, 
                         original_doc: Any = None, is_docx: bool = None) -> Tuple[Optional[str], Optional[str], Optional[str]]:
    """Save content to a document file with appropriate formatting.
    
    Args:
        content: Text content to write.
        file_path: Path where to save the document.
        original_doc: Original document object to preserve formatting.
        is_docx: Boolean indicating if the file is a DOCX (True) or ODT (False).
        
    Returns:
        Tuple of (json_path, text_path, doc_path)
    """
    try:
        # Create a correction entry for the full text
        corrections = {
            'corrections': [{
                'original': '',  # Empty since this is a full text replacement
                'corrected': content,
                'type': 'full_text'
            }],
            'corrected_full_text': content
        }
        
        return save_document(file_path, corrections, original_doc, is_docx)
    except Exception as e:
        print(f"Error saving document content: {e}")
        return None, None, None

def save_correction_outputs(file_path: str, corrections: List[Dict[str, Any]], 
                           original_doc: Any = None, is_docx: bool = True) -> Tuple[Optional[str], Optional[str], Optional[str]]:
    """Save correction outputs to files.
    This function is maintained for compatibility but delegates to save_document.
    
    Args:
        file_path: Original document file path
        corrections: List of corrections or dictionary with corrections
        original_doc: Original document object
        is_docx: Whether the document is a DOCX file
        
    Returns:
        Tuple of (json_path, text_path, doc_path)
    """
    # Just delegate to the main save_document function
    return save_document(file_path, corrections, original_doc, is_docx)

def create_fallback_document(corrections, output_path, is_docx):
    """Create a simple document with corrections when main save fails.
    
    Args:
        corrections: Correction data (dict or list)
        output_path: Output file path
        is_docx: Whether to create DOCX (True) or ODT (False)
        
    Returns:
        Path to the saved document
    """
    try:
        # Convert corrections to list format if needed
        if isinstance(corrections, dict):
            corrections_list = corrections.get('corrections', [])
        elif isinstance(corrections, list):
            corrections_list = corrections
        else:
            corrections_list = [corrections] if corrections else []
        
        # Create text file with corrections
        text_path = f"{output_path}.txt"
        with open(text_path, 'w', encoding='utf-8') as f:
            f.write("CORRECTIONS:\n\n")
            for i, corr in enumerate(corrections_list, 1):
                if isinstance(corr, dict):
                    f.write(f"{i}. Original: {corr.get('original', '')}\n")
                    f.write(f"   Corrected: {corr.get('corrected', '')}\n")
                    if 'explanation' in corr:
                        f.write(f"   Explanation: {corr['explanation']}\n")
                elif isinstance(corr, str):
                    f.write(f"{i}. Correction: {corr}\n")
                f.write("\n")
        
        # Try to create a simple document if possible
        if is_docx:
            try:
                from docx import Document
                doc = Document()
                for i, corr in enumerate(corrections_list, 1):
                    if isinstance(corr, dict):
                        p = doc.add_paragraph(f"{i}. Original: ")
                        p.add_run(corr.get('original', '')).bold = True
                        p.add_run("\nCorrected: ")
                        p.add_run(corr.get('corrected', '')).italic = True
                        if 'explanation' in corr:
                            p.add_run("\nExplanation: ")
                            p.add_run(corr.get('explanation', ''))
                    else:
                        doc.add_paragraph(f"{i}. {corr}")
                    doc.add_paragraph("")
                
                doc_path = f"{output_path}.docx"
                doc.save(doc_path)
                return doc_path
            except Exception as e:
                print(f"Could not create fallback DOCX: {e}")
                
        return text_path
        
    except Exception as e:
        print(f"Error creating fallback document: {e}")
        return None

def create_sample_document(file_path):
    """Create a sample document file for testing.
    
    Args:
        file_path: Path where the sample document will be saved
    """
    # Check the file extension
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.docx':
        try:
            from docx import Document
            from docx.shared import Pt
            
            doc = Document()
            doc.add_heading('Sample Document', 0)
            
            p = doc.add_paragraph('This is a sample document created for testing. ')
            p.add_run('It contains some text with potential grammar issues that can be corrected.')
            
            doc.add_heading('Section with Grammar Issues', level=1)
            
            doc.add_paragraph('This sentence have a verb agreement error.')
            doc.add_paragraph('There are mistakes in this sentance.')
            doc.add_paragraph('This paragraph contains multiple erors that should be corected by the grammar checker.')
            
            doc.add_heading('Section with Formatting', level=1)
            p = doc.add_paragraph('This section has ')
            p.add_run('bold').bold = True
            p.add_run(' and ')
            p.add_run('italic').italic = True
            p.add_run(' text formatting.')
            
            # Save the document
            doc.save(file_path)
            print(f"Created sample DOCX file at {file_path}")
            
        except Exception as e:
            print(f"Error creating sample DOCX: {e}")
            
    elif file_ext == '.odt':
        try:
            # For ODT, we'll create an ODT file using odfpy library
            try:
                from odf.opendocument import OpenDocumentText
                from odf.style import Style, TextProperties, ParagraphProperties
                from odf.text import H, P, Span
                
                textdoc = OpenDocumentText()
                
                # Add styles
                heading_style = Style(name="Heading", family="paragraph")
                heading_style.addElement(TextProperties(fontsize="16pt", fontweight="bold"))
                textdoc.styles.addElement(heading_style)
                
                # Add content
                h = H(outlinelevel=1, text="Sample ODT Document")
                textdoc.text.addElement(h)
                
                p = P(text="This is a sample document created for testing. It contains some text with potential grammar issues that can be corrected.")
                textdoc.text.addElement(p)
                
                h = H(outlinelevel=2, text="Section with Grammar Issues")
                textdoc.text.addElement(h)
                
                p = P(text="This sentence have a verb agreement error.")
                textdoc.text.addElement(p)
                
                p = P(text="There are mistakes in this sentance.")
                textdoc.text.addElement(p)
                
                p = P(text="This paragraph contains multiple erors that should be corected by the grammar checker.")
                textdoc.text.addElement(p)
                
                textdoc.save(file_path)
                print(f"Created sample ODT file at {file_path}")
                return
            except ImportError:
                print("odfpy library not found, using alternative method")
            
            # Alternative method using a template and text replacement
            import zipfile
            import tempfile
            import shutil
            import os
            
            # Create a simple ODT file structure
            with tempfile.TemporaryDirectory() as temp_dir:
                # Create META-INF directory and mimetype
                os.makedirs(os.path.join(temp_dir, 'META-INF'))
                
                # Create content.xml
                with open(os.path.join(temp_dir, 'content.xml'), 'w', encoding='utf-8') as f:
                    f.write('''<?xml version="1.0" encoding="UTF-8"?>
<office:document-content xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" 
                         xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" 
                         xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" 
                         xmlns:table="urn:oasis:names:tc:opendocument:xmlns:table:1.0" 
                         xmlns:draw="urn:oasis:names:tc:opendocument:xmlns:drawing:1.0" 
                         xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" 
                         xmlns:xlink="http://www.w3.org/1999/xlink" 
                         xmlns:dc="http://purl.org/dc/elements/1.1/" 
                         xmlns:meta="urn:oasis:names:tc:opendocument:xmlns:meta:1.0" 
                         xmlns:number="urn:oasis:names:tc:opendocument:xmlns:datastyle:1.0" 
                         xmlns:svg="urn:oasis:names:tc:opendocument:xmlns:svg-compatible:1.0" 
                         xmlns:chart="urn:oasis:names:tc:opendocument:xmlns:chart:1.0" 
                         xmlns:dr3d="urn:oasis:names:tc:opendocument:xmlns:dr3d:1.0" 
                         xmlns:math="http://www.w3.org/1998/Math/MathML" 
                         xmlns:form="urn:oasis:names:tc:opendocument:xmlns:form:1.0" 
                         xmlns:script="urn:oasis:names:tc:opendocument:xmlns:script:1.0" 
                         xmlns:ooo="http://openoffice.org/2004/office" 
                         xmlns:ooow="http://openoffice.org/2004/writer" 
                         xmlns:oooc="http://openoffice.org/2004/calc" 
                         xmlns:dom="http://www.w3.org/2001/xml-events" 
                         xmlns:xforms="http://www.w3.org/2002/xforms" 
                         xmlns:xsd="http://www.w3.org/2001/XMLSchema" 
                         xmlns:xsi="http://www.w3.org/2001/XMLSchema-instance" 
                         office:version="1.2">
  <office:body>
    <office:text>
      <text:h text:style-name="Heading_20_1" text:outline-level="1">Sample ODT Document</text:h>
      <text:p text:style-name="Text_20_body">This is a sample document created for testing. It contains some text with potential grammar issues that can be corrected.</text:p>
      <text:h text:style-name="Heading_20_2" text:outline-level="2">Section with Grammar Issues</text:h>
      <text:p text:style-name="Text_20_body">This sentence have a verb agreement error.</text:p>
      <text:p text:style-name="Text_20_body">There are mistakes in this sentance.</text:p>
      <text:p text:style-name="Text_20_body">This paragraph contains multiple erors that should be corected by the grammar checker.</text:p>
    </office:text>
  </office:body>
</office:document-content>''')
                
                # Create mimetype file
                with open(os.path.join(temp_dir, 'mimetype'), 'w', encoding='utf-8') as f:
                    f.write('application/vnd.oasis.opendocument.text')
                
                # Create META-INF/manifest.xml
                with open(os.path.join(temp_dir, 'META-INF', 'manifest.xml'), 'w', encoding='utf-8') as f:
                    f.write('''<?xml version="1.0" encoding="UTF-8"?>
<manifest:manifest xmlns:manifest="urn:oasis:names:tc:opendocument:xmlns:manifest:1.0">
 <manifest:file-entry manifest:media-type="application/vnd.oasis.opendocument.text" manifest:full-path="/"/>
 <manifest:file-entry manifest:media-type="text/xml" manifest:full-path="content.xml"/>
 <manifest:file-entry manifest:media-type="text/xml" manifest:full-path="META-INF/manifest.xml"/>
</manifest:manifest>''')
                
                # Create styles.xml
                with open(os.path.join(temp_dir, 'styles.xml'), 'w', encoding='utf-8') as f:
                    f.write('''<?xml version="1.0" encoding="UTF-8"?>
<office:document-styles xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" 
                        xmlns:style="urn:oasis:names:tc:opendocument:xmlns:style:1.0" 
                        xmlns:text="urn:oasis:names:tc:opendocument:xmlns:text:1.0" 
                        xmlns:fo="urn:oasis:names:tc:opendocument:xmlns:xsl-fo-compatible:1.0" 
                        office:version="1.2">
  <office:styles>
    <style:style style:name="Standard" style:family="paragraph" style:class="text"/>
    <style:style style:name="Heading" style:family="paragraph" style:parent-style-name="Standard" style:class="text">
      <style:text-properties fo:font-size="14pt" fo:font-weight="bold"/>
    </style:style>
    <style:style style:name="Text_20_body" style:display-name="Text body" style:family="paragraph" style:parent-style-name="Standard" style:class="text">
      <style:paragraph-properties fo:margin-top="0cm" fo:margin-bottom="0.212cm"/>
    </style:style>
    <style:style style:name="Heading_20_1" style:display-name="Heading 1" style:family="paragraph" style:parent-style-name="Heading" style:next-style-name="Text_20_body" style:class="text">
      <style:text-properties fo:font-size="18pt" fo:font-weight="bold"/>
    </style:style>
    <style:style style:name="Heading_20_2" style:display-name="Heading 2" style:family="paragraph" style:parent-style-name="Heading" style:next-style-name="Text_20_body" style:class="text">
      <style:text-properties fo:font-size="16pt" fo:font-weight="bold"/>
    </style:style>
  </office:styles>
</office:document-styles>''')
                
                # Create meta.xml
                with open(os.path.join(temp_dir, 'meta.xml'), 'w', encoding='utf-8') as f:
                    f.write('''<?xml version="1.0" encoding="UTF-8"?>
<office:document-meta xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" 
                     xmlns:meta="urn:oasis:names:tc:opendocument:xmlns:meta:1.0" 
                     xmlns:dc="http://purl.org/dc/elements/1.1/" 
                     office:version="1.2">
  <office:meta>
    <dc:title>Sample ODT Document</dc:title>
    <dc:creator>Chaiotic Grammar Checker</dc:creator>
    <dc:date>2023-01-01T00:00:00</dc:date>
  </office:meta>
</office:document-meta>''')
                
                # Create settings.xml
                with open(os.path.join(temp_dir, 'settings.xml'), 'w', encoding='utf-8') as f:
                    f.write('''<?xml version="1.0" encoding="UTF-8"?>
<office:document-settings xmlns:office="urn:oasis:names:tc:opendocument:xmlns:office:1.0" 
                         xmlns:config="urn:oasis:names:tc:opendocument:xmlns:config:1.0" 
                         office:version="1.2">
  <office:settings>
    <config:config-item-set config:name="ooo:view-settings"/>
  </office:settings>
</office:document-settings>''')
                
                # Create the ODT file (which is a ZIP file)
                with zipfile.ZipFile(file_path, 'w') as zip_file:
                    # Add mimetype first without compression
                    zip_file.write(os.path.join(temp_dir, 'mimetype'), 'mimetype', compress_type=zipfile.ZIP_STORED)
                    
                    # Add other files with compression
                    for root, dirs, files in os.walk(temp_dir):
                        for file in files:
                            if file != 'mimetype':  # Skip mimetype as we already added it
                                file_path_in_zip = os.path.relpath(os.path.join(root, file), temp_dir)
                                zip_file.write(os.path.join(root, file), file_path_in_zip, compress_type=zipfile.ZIP_DEFLATED)
                
                print(f"Created sample ODT file at {file_path}")
                
        except Exception as e:
            print(f"Error creating sample ODT: {e}")
            import traceback
            traceback.print_exc()
    else:
        print(f"Cannot create sample document with extension {file_ext}. Please use .docx or .odt.")

def extract_structured_content(file_path):
    """Extract structured content from a document file.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        List of dictionaries containing structured content
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    # Call the specific extractor function based on file type
    if file_ext == '.odt':
        from .document_extractor import extract_structured_odt
        content = extract_structured_odt(file_path)
    elif file_ext == '.docx':
        from .document_extractor import extract_structured_docx
        content = extract_structured_docx(file_path)
    elif file_ext == '.txt':
        from .document_extractor import extract_structured_text
        content = extract_structured_text(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")
    
    # Debug output - show what was extracted
    print(f"\n===== EXTRACTED DOCUMENT STRUCTURE ({len(content)} elements) =====")
    for i, item in enumerate(content[:3], 1):  # Print first 3 items for preview
        print(f"\nItem {i}/{len(content)}: ID={item.get('id')}, Type={item.get('type')}")
        print(f"Content: {item.get('content')[:100]}{'...' if len(item.get('content', '')) > 100 else ''}")
        
        # Check if we have XML content
        if 'xml_content' in item:
            xml_preview = item['xml_content'][:150]
            print(f"XML: {xml_preview}{'...' if len(item['xml_content']) > 150 else ''}")
        
        # Check for nested elements
        if 'metadata' in item and 'nested_elements' in item['metadata']:
            nested = item['metadata']['nested_elements']
            print(f"Nested elements: {len(nested)}")
            if nested:
                print(f"  First nested: {nested[0].get('type')}, Text: {nested[0].get('text')[:50]}...")
    
    if len(content) > 3:
        print(f"\n... and {len(content) - 3} more items\n")
    
    print("=" * 60)
    
    return content

# Re-export all the functions to maintain backward compatibility
__all__ = [
    'read_document',
    'save_document',
    'save_document_content',
    'save_correction_outputs',
    'create_sample_document',
    'extract_structured_content'
]