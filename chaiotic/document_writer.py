"""Module for saving documents with corrections."""

import os
from typing import List, Dict, Any

# Changed to import from utils module at root level
from utils.text_utils import FuzzyMatcher, generate_full_text_from_corrections
from utils.xml_utils import (create_xml_element, create_tracked_change_region, 
                      parse_xml_file, write_xml_file, LXML_AVAILABLE)

def save_document(file_path: str, corrections: List[Dict[Any, Any]], original_doc=None, is_docx: bool = True) -> str:
    """
    Save the document with corrections as Word comments or ODT tracked changes.
    This function preserves the original document and adds review comments/changes.
    
    Args:
        file_path: Path to the original document
        corrections: List of corrections to apply
        original_doc: Original document object if available
        is_docx: Whether the document is DOCX (otherwise ODT)
        
    Returns:
        Path to the saved document
    """
    # Generate output filename
    file_base, file_ext = os.path.splitext(file_path)
    output_path = f"{file_base}_corrected{file_ext}"
    
    print(f"Saving document with corrections to {output_path}")
    print(f"Number of corrections to apply: {len(corrections)}")
    
    # Create an instance of our optimized fuzzy matcher
    matcher = FuzzyMatcher()
    
    # Validate corrections to ensure they are all properly formatted dictionaries
    valid_corrections = []
    for correction in corrections:
        if not isinstance(correction, dict):
            print(f"Warning: Skipping non-dictionary correction: {correction}")
            continue
        if 'original' not in correction or 'corrected' not in correction:
            print(f"Warning: Skipping correction missing required fields: {correction}")
            continue
        valid_corrections.append(correction)
    
    print(f"After validation: {len(valid_corrections)} valid corrections")
    corrections = valid_corrections
    
    if is_docx:
        # If we don't have the original document, load it now
        if original_doc is None:
            from docx import Document
            original_doc = Document(file_path)
        
        # Apply corrections directly to a copy of the document
        corrected_doc = apply_corrections_to_document(original_doc, corrections)
        
        # Save the corrected document
        try:
            corrected_doc.save(output_path)
            print(f"Successfully saved corrected document to {output_path}")
            return output_path
        except Exception as e:
            print(f"Error saving document: {str(e)}")
            # Try saving to a different location
            fallback_path = os.path.join(os.path.dirname(file_path), "corrected_output.docx")
            try:
                corrected_doc.save(fallback_path)
                print(f"Saved to alternate location: {fallback_path}")
                return fallback_path
            except Exception as e2:
                print(f"Failed to save document: {str(e2)}")
                raise
    else:
        # Handle ODT files with tracked changes
        print("Processing ODT file with tracked changes...")
        output_path = apply_corrections_to_odt(file_path, corrections, output_path)
        return output_path

def save_correction_outputs(file_path, corrections, original_doc, is_docx=True):
    """Save correction outputs to files.
    
    Args:
        file_path: Original document file path
        corrections: List of corrections or dictionary with corrections
        original_doc: Original document object
        is_docx: Whether the document is a DOCX file
        
    Returns:
        Tuple of (json_path, text_path, doc_path)
    """
    import os
    import json
    import shutil
    
    # Generate output file paths
    base_name = os.path.splitext(file_path)[0]
    json_path = f"{base_name}_corrections.json"
    text_path = f"{base_name}_corrections.txt"
    doc_path = f"{base_name}_corrected{'.docx' if is_docx else '.odt'}"
    
    # Process corrections into a consistent format
    if isinstance(corrections, dict):
        # Handle dictionary format (either with 'corrections' key or direct corrections)
        if 'corrections' in corrections:
            corrections_list = corrections['corrections']
            corrected_full_text = corrections.get('corrected_full_text', '')
        else:
            # It's a direct dictionary of corrections
            corrections_list = [corrections]
            corrected_full_text = ''
    elif isinstance(corrections, list):
        # Already a list of corrections
        corrections_list = corrections
        corrected_full_text = ''
    else:
        # Unknown format
        print(f"Warning: Unknown corrections format: {type(corrections)}")
        corrections_list = []
        corrected_full_text = ''
    
    # Ensure all items in corrections_list are dictionaries
    valid_corrections = []
    for item in corrections_list:
        if isinstance(item, dict) and 'original' in item and 'corrected' in item:
            valid_corrections.append(item)
        else:
            print(f"Warning: Skipping invalid correction: {item}")
    
    corrections_list = valid_corrections
    print(f"Number of valid corrections to process: {len(corrections_list)}")
    
    # Generate full text from corrections if not provided
    if not corrected_full_text:
        corrected_full_text = generate_full_text_from_corrections(corrections_list)
    
    # Save JSON file with corrections
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump({
            'corrections': corrections_list,
            'corrected_full_text': corrected_full_text
        }, f, indent=2, ensure_ascii=False)
    
    # Save text file with corrected content
    with open(text_path, 'w', encoding='utf-8') as f:
        f.write(corrected_full_text)
    
    # Save corrected document file
    try:
        if is_docx:
            # Use the apply_corrections_to_document function from this module
            corrected_doc = apply_corrections_to_document(original_doc, corrections_list)
            corrected_doc.save(doc_path)
        else:
            # For ODT, apply corrections with tracked changes
            apply_corrections_to_odt(file_path, corrections_list, doc_path)
    except Exception as e:
        print(f"Error saving corrected document: {e}")
        print("Falling back to basic text document creation...")
        try:
            # Simple fallback to create a basic document
            if is_docx:
                from docx import Document
                doc = Document()
                for line in corrected_full_text.split('\n'):
                    doc.add_paragraph(line)
                doc.save(doc_path)
            else:
                # For ODT, just copy the text file
                shutil.copy(text_path, doc_path)
        except Exception as e2:
            print(f"Error creating basic document: {e2}")
    
    return json_path, text_path, doc_path

def apply_corrections_to_odt(file_path, corrections, output_path):
    """Apply corrections to an ODT document with tracked changes."""
    import zipfile
    import tempfile
    import xml.etree.ElementTree as ET
    from datetime import datetime
    
    # Use the xml_utils module functions if available
    if LXML_AVAILABLE:
        print("Using lxml for better XML handling")
        import lxml.etree as LET
        
        def apply_tracked_changes(root, corrections, user, timestamp):
            """Apply tracked changes to the document content."""
            nsmap = {
                'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0',
                'office': 'urn:oasis:names:tc:opendocument:xmlns:office:1.0',
                'dc': 'http://purl.org/dc/elements/1.1/'
            }
            
            # Find or create tracked-changes element
            office_text = root.find('.//office:text', nsmap)
            if office_text is None:
                raise ValueError("Could not find office:text element")
                
            tracked_changes = office_text.find('.//text:tracked-changes', nsmap)
            if tracked_changes is None:
                tracked_changes = create_xml_element('text:tracked-changes', nsmap=nsmap)
                office_text.insert(0, tracked_changes)
            
            # Process paragraphs
            paragraphs = root.findall('.//text:p', nsmap)
            changes_by_para = {}
            
            # Initialize a master change_id at module level to ensure unique IDs across all changes
            master_change_id = 1
            
            # First pass: Identify changes and create change regions
            for i, para in enumerate(paragraphs):
                para_text = "".join(para.xpath('.//text()'))
                para_id = f"p{i+1}"
                
                para_changes = []
                for corr in corrections:
                    original = corr.get('original', '')
                    if original and original in para_text:
                        # Using the master_change_id to ensure unique IDs
                        del_change_id = f"ctd{master_change_id}"
                        ins_change_id = f"cti{master_change_id}"
                        master_change_id += 1  # Increment for next change
                        
                        change = {
                            'del_id': del_change_id,
                            'ins_id': ins_change_id,
                            'original': original,
                            'corrected': corr.get('corrected', ''),
                            'start': para_text.find(original),
                            'end': para_text.find(original) + len(original)
                        }
                        para_changes.append(change)
                        
                        # Create two separate regions - one for deletion and one for insertion
                        # 1. Create deletion region
                        deletion_region = create_tracked_change_region(
                            del_change_id, 'deletion', original, user, timestamp, nsmap
                        )
                        tracked_changes.append(deletion_region)
                        
                        # 2. Create insertion region
                        insertion_region = create_tracked_change_region(
                            ins_change_id, 'insertion', corr.get('corrected', ''), 
                            user, timestamp, nsmap
                        )
                        tracked_changes.append(insertion_region)
                        
                if para_changes:
                    changes_by_para[para_id] = sorted(para_changes, key=lambda x: x['start'])
            
            # Second pass: Apply changes to paragraphs
            for para_id, changes in changes_by_para.items():
                para = paragraphs[int(para_id[1:]) - 1]
                para_text = "".join(para.xpath('.//text()'))
                
                # Clear existing content
                for child in para:
                    para.remove(child)
                para.text = None
                
                last_pos = 0
                for change in changes:
                    # Add text before change
                    if change['start'] > last_pos:
                        before_text = para_text[last_pos:change['start']]
                        if last_pos == 0:
                            para.text = before_text
                        else:
                            span = create_xml_element('text:span', text=before_text, nsmap=nsmap)
                            para.append(span)
                    
                    # Add deletion marker
                    change_start = create_xml_element('text:change-start', 
                                                    {'text:change-id': change['del_id']}, 
                                                    nsmap=nsmap)
                    para.append(change_start)
                    
                    # Add change end for deletion
                    change_end = create_xml_element('text:change-end',
                                                  {'text:change-id': change['del_id']},
                                                  nsmap=nsmap)
                    para.append(change_end)
                    
                    # Add insertion marker
                    change_start = create_xml_element('text:change-start', 
                                                    {'text:change-id': change['ins_id']}, 
                                                    nsmap=nsmap)
                    para.append(change_start)
                    
                    # Add corrected text
                    span = create_xml_element('text:span', text=change['corrected'], nsmap=nsmap)
                    para.append(span)
                    
                    # Add change end for insertion
                    change_end = create_xml_element('text:change-end',
                                                  {'text:change-id': change['ins_id']},
                                                  nsmap=nsmap)
                    para.append(change_end)
                    
                    last_pos = change['end']
                
                # Add remaining text
                if last_pos < len(para_text):
                    span = create_xml_element('text:span', text=para_text[last_pos:], nsmap=nsmap)
                    para.append(span)
            
            return root
        
        # Main ODT processing 
        with tempfile.TemporaryDirectory() as temp_dir:
            with zipfile.ZipFile(file_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            # Parse and modify content.xml
            content_file = os.path.join(temp_dir, 'content.xml')
            parser = LET.XMLParser(remove_blank_text=True)
            tree = parse_xml_file(content_file, parser)
            
            # Apply changes
            timestamp = datetime.now().strftime("%Y-%m-%dT%H:%M:%S")
            modified_root = apply_tracked_changes(tree.getroot(), corrections, 
                                              "Chaiotic Grammar Checker", timestamp)
            
            # Write modified content back
            write_xml_file(tree, content_file)
            
            # Create new ODT file
            with zipfile.ZipFile(output_path, 'w') as new_odt:
                # Add mimetype first without compression
                mimetype_path = os.path.join(temp_dir, 'mimetype')
                if os.path.exists(mimetype_path):
                    new_odt.write(mimetype_path, 'mimetype', compress_type=zipfile.ZIP_STORED)
                
                # Add all other files with compression
                for root, _, files in os.walk(temp_dir):
                    for file in files:
                        if file != 'mimetype':
                            file_path = os.path.join(root, file)
                            arcname = os.path.relpath(file_path, temp_dir)
                            new_odt.write(file_path, arcname, compress_type=zipfile.ZIP_DEFLATED)
            
            return output_path
    else:
        print("lxml not available, falling back to basic corrections")
        return apply_corrections_to_odt_basic(file_path, corrections, output_path)

def apply_corrections_to_odt_basic(file_path, corrections, output_path):
    """Basic fallback for ODT corrections when lxml is not available."""
    # Create a simple text file with corrections
    with open(output_path + '.txt', 'w', encoding='utf-8') as f:
        f.write("CORRECTION SUGGESTIONS:\n\n")
        for i, correction in enumerate(corrections, 1):
            f.write(f"{i}. Original: {correction.get('original', 'N/A')}\n")
            f.write(f"   Corrected: {correction.get('corrected', 'N/A')}\n")
            f.write(f"   Explanation: {correction.get('explanation', 'N/A')}\n\n")
    
    return output_path + '.txt'

def apply_corrections_to_document(original_doc, corrections, is_docx=True):
    """Apply corrections to a document as tracked changes suggestions.
    
    Args:
        original_doc: Original document object
        corrections: List of corrections to apply
        is_docx: Whether the document is a DOCX file
        
    Returns:
        Corrected document object with tracked changes
    """
    if is_docx:
        from docx import Document
        from docx.shared import RGBColor
        
        # Always create a new document - this is the most reliable approach
        corrected_doc = Document()
        
        # Copy document styles from the original document to the new document
        try:
            # Copy the base styles
            for style_id in original_doc.styles:
                if style_id not in corrected_doc.styles:
                    try:
                        style = original_doc.styles[style_id]
                        # We can't directly copy styles between documents,
                        # but we can ensure the style exists in the target document
                        if style.type == 1:  # Paragraph style
                            if style.name not in corrected_doc.styles:
                                corrected_doc.styles.add_style(style.name, style.type)
                    except Exception as e:
                        print(f"Warning: Could not copy style {style_id}: {e}")
        except Exception as e:
            print(f"Warning: Error copying document styles: {e}")
        
        # Copy the document properties and sections
        for section in original_doc.sections:
            new_section = corrected_doc.add_section()
            try:
                new_section.page_height = section.page_height
                new_section.page_width = section.page_width
                new_section.left_margin = section.left_margin
                new_section.right_margin = section.right_margin
                new_section.top_margin = section.top_margin
                new_section.bottom_margin = section.bottom_margin
                new_section.header_distance = section.header_distance
                new_section.footer_distance = section.footer_distance
            except Exception as e:
                print(f"Warning: Could not copy section properties: {e}")
        
        # Create a mapping of corrections by paragraph ID
        correction_by_paragraph = {}
        for correction in corrections:
            if 'id' in correction and correction['id'].startswith('p'):
                para_id = correction['id']
                if para_id not in correction_by_paragraph:
                    correction_by_paragraph[para_id] = []
                correction_by_paragraph[para_id].append(correction)
        
        # Debug: Print corrections mapped to paragraphs
        print(f"Corrections mapped by paragraph ID: {len(correction_by_paragraph)}")
        
        # Apply corrections to paragraphs
        for i, original_para in enumerate(original_doc.paragraphs):
            para_id = f"p{i+1}"
            
            # Create a new paragraph in the corrected document
            new_para = corrected_doc.add_paragraph()
            
            # Copy the paragraph style
            if hasattr(original_para, 'style') and original_para.style:
                try:
                    style_name = original_para.style.name
                    if style_name in corrected_doc.styles:
                        new_para.style = style_name
                    else:
                        try:
                            corrected_doc.styles.add_style(style_name, 1)  # 1 = paragraph style
                            new_para.style = style_name
                        except:
                            new_para.style = 'Normal'
                except Exception as e:
                    print(f"Warning: Could not copy style for paragraph {para_id}: {e}")
            
            # Check if we have corrections for this paragraph
            para_corrections = correction_by_paragraph.get(para_id, [])
            
            if para_corrections and original_para.text.strip():
                # We have corrections for this paragraph
                print(f"Applying corrections to paragraph {para_id}: {original_para.text}")
                
                # Get the original text
                original_text = original_para.text
                remaining_text = original_text
                
                # Track if we've applied any corrections
                applied_corrections = False
                
                # Apply each correction
                for correction in para_corrections:
                    if 'original' in correction and 'corrected' in correction:
                        original_part = correction.get('original', '')
                        corrected_part = correction.get('corrected', '')
                        explanation = correction.get('explanation', '')
                        
                        # Debug
                        print(f"Looking for: '{original_part}' to replace with '{corrected_part}'")
                        print(f"In text: '{remaining_text}'")
                        
                        if original_part in remaining_text:
                            # Split the text at the correction point
                            before, match, after = remaining_text.partition(original_part)
                            
                            # Add the text before the correction
                            if before:
                                new_para.add_run(before)
                            
                            # Add the correction with visual formatting
                            # Original with strikethrough
                            del_run = new_para.add_run(original_part)
                            del_run.font.strike = True
                            
                            # Corrected in red
                            ins_run = new_para.add_run(corrected_part)
                            ins_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                            ins_run.bold = True
                                                        
                            # Update remaining text for next correction
                            remaining_text = after
                            applied_corrections = True
                            print(f"Applied correction: '{original_part}' -> '{corrected_part}'")
                        else:
                            print(f"WARNING: Could not find text '{original_part}' in remaining paragraph.")
                
                # Add any remaining text
                if remaining_text:
                    new_para.add_run(remaining_text)
                
                # If no corrections were applied, just add the original text
                if not applied_corrections:
                    if len(new_para.runs) == 0:  # Ensure we haven't already added content
                        new_para.text = original_text
                    print(f"No corrections were successfully applied to paragraph {para_id}")
            else:
                # No corrections for this paragraph, just copy the text and formatting
                if original_para.runs:
                    # Copy runs to preserve formatting
                    for run in original_para.runs:
                        new_run = new_para.add_run(run.text)
                        # Simple formatting copy
                        try:
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            if hasattr(run, 'font') and hasattr(run.font, 'color') and run.font.color:
                                new_run.font.color.rgb = run.font.color.rgb
                        except Exception as e:
                            print(f"Warning: Could not copy run formatting: {e}")
                else:
                    # If no runs, just copy the text
                    new_para.text = original_para.text
        
        # Copy tables
        for table in original_doc.tables:
            try:
                new_table = corrected_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                # Copy table style if available
                if hasattr(table, 'style') and table.style:
                    try:
                        new_table.style = table.style
                    except:
                        print("Warning: Could not copy table style")
                
                # Copy cell contents
                for i, row in enumerate(table.rows):
                    for j, cell in enumerate(row.cells):
                        try:
                            # Copy each paragraph in the cell
                            for para in cell.paragraphs:
                                cell_para = new_table.rows[i].cells[j].add_paragraph()
                                cell_para.text = para.text
                                # Try to copy paragraph style
                                if hasattr(para, 'style') and para.style:
                                    try:
                                        cell_para.style = para.style.name
                                    except:
                                        pass
                        except Exception as e:
                            # Fall back to simple text copy
                            new_table.rows[i].cells[j].text = cell.text
                            print(f"Warning: Error copying table cell content: {e}")
            except Exception as e:
                print(f"Warning: Could not copy table: {e}")
        
        return corrected_doc
    
    else:
        # For ODT files, we don't have direct editing capability
        return original_doc