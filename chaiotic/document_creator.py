"""Module for creating sample documents for testing."""

import os
import re
from typing import Optional

def create_sample_document(file_path, content=None):
    """Create a sample document for testing.
    
    Args:
        file_path: Path where to save the document.
        content: Optional content to include in the document.
        
    Returns:
        Path to the created document.
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if not content:
        content = (
            "Überschrift 1\n\n"
            "Dies ist ein Absatz mit einigen falsch geschriebenen Wörtern. "
            "Zum Beispeil hier und drot und villeicht auch anderswo.\n\n"
            "Überschrift 2\n\n"
            "Dieser Abschnit enthält auch Grammatikfehler, die der korrektur bedürfen. "
            "Die Katze hast auf dem Sofa gesessen, obwohl sie dass nicht tun sollte.\n\n"
            "Manchmal fehlen auch Kommas oder es gibt zu viele davon, oder die Satzstellung ist falsch verwendet worden."
        )
    
    if file_ext == '.docx':
        return create_sample_docx(file_path, content)
    elif file_ext == '.odt':
        return create_sample_odt(file_path, content)
    elif file_ext == '.txt':
        return create_sample_txt(file_path, content)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")

def create_sample_docx(file_path, content):
    """Create a sample DOCX document."""
    try:
        from docx import Document
        doc = Document()
        
        # Split into paragraphs
        paragraphs = re.split(r'\n\s*\n', content)
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Check if it looks like a heading (short, single line)
            lines = para.split('\n')
            if len(lines) == 1 and len(para) < 100 and not para.endswith('.'):
                # Probable heading - add as Heading 1
                doc.add_heading(para, level=1)
            else:
                # Regular paragraph
                doc.add_paragraph(para)
        
        doc.save(file_path)
        print(f"Sample DOCX created at {file_path}")
        return file_path
    except ImportError:
        print("python-docx library not available. Creating a plain text file instead.")
        return create_sample_txt(os.path.splitext(file_path)[0] + '.txt', content)
    except Exception as e:
        print(f"Error creating sample DOCX: {e}")
        # Fallback to text file
        return create_sample_txt(os.path.splitext(file_path)[0] + '.txt', content)

def create_sample_odt(file_path, content):
    """Create a sample ODT document."""
    try:
        from odf.opendocument import OpenDocumentText
        from odf.style import Style, TextProperties, ParagraphProperties
        from odf.text import H, P
        
        doc = OpenDocumentText()
        
        # Create heading style
        heading_style = Style(name="Heading", family="paragraph")
        heading_style.addElement(TextProperties(fontweight="bold", fontsize="16pt"))
        doc.styles.addElement(heading_style)
        
        # Split into paragraphs
        paragraphs = re.split(r'\n\s*\n', content)
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Check if it looks like a heading (short, single line)
            lines = para.split('\n')
            if len(lines) == 1 and len(para) < 100 and not para.endswith('.'):
                # Probable heading - add as Heading
                heading = H(stylename=heading_style, outlinelevel=1)
                heading.addText(para)
                doc.text.addElement(heading)
            else:
                # Regular paragraph
                p = P()
                p.addText(para)
                doc.text.addElement(p)
        
        doc.save(file_path)
        print(f"Sample ODT created at {file_path}")
        return file_path
    except ImportError:
        print("odfpy library not available. Creating a plain text file instead.")
        return create_sample_txt(os.path.splitext(file_path)[0] + '.txt', content)
    except Exception as e:
        print(f"Error creating sample ODT: {e}")
        # Fallback to text file
        return create_sample_txt(os.path.splitext(file_path)[0] + '.txt', content)

def create_sample_txt(file_path, content):
    """Create a sample text document."""
    try:
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"Sample text file created at {file_path}")
        return file_path
    except Exception as e:
        print(f"Error creating sample text file: {e}")
        raise