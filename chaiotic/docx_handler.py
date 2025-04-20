"""Module for handling DOCX document operations."""

import os
from typing import List, Dict, Any, Optional, Union

def apply_corrections_to_docx(file_path: str, corrections: List[Dict[str, Any]], 
                             output_path: str) -> str:
    """Apply corrections to a DOCX document with tracked changes.
    
    Args:
        file_path: Path to the original DOCX document
        corrections: List of correction dictionaries
        output_path: Path to save the corrected document
        
    Returns:
        Path to the saved document
    """
    try:
        from docx import Document
        from docx.shared import RGBColor
        
        # Load the original document
        original_doc = Document(file_path)
        
        # Always create a new document - this is most reliable approach
        corrected_doc = Document()
        
        # Copy document styles
        copy_docx_formatting(original_doc, corrected_doc)
        
        # Create a mapping of corrections by paragraph ID
        correction_by_paragraph = {}
        for correction in corrections:
            if 'id' in correction and correction['id'].startswith('p'):
                para_id = correction['id']
                if para_id not in correction_by_paragraph:
                    correction_by_paragraph[para_id] = []
                correction_by_paragraph[para_id].append(correction)
        
        # Apply corrections to paragraphs
        for i, original_para in enumerate(original_doc.paragraphs):
            para_id = f"p{i+1}"
            
            # Create a new paragraph in the corrected document
            new_para = corrected_doc.add_paragraph()
            
            # Copy the paragraph style
            if hasattr(original_para, 'style') and original_para.style:
                try:
                    style_name = original_para.style.name
                    if style_name in corrected_doc.styles:
                        new_para.style = style_name
                    else:
                        try:
                            corrected_doc.styles.add_style(style_name, 1)  # 1 = paragraph style
                            new_para.style = style_name
                        except:
                            new_para.style = 'Normal'
                except Exception as e:
                    print(f"Warning: Could not copy style for paragraph {para_id}: {e}")
            
            # Check if we have corrections for this paragraph
            para_corrections = correction_by_paragraph.get(para_id, [])
            has_correction = False
            
            # Get list of corrections without paragraph ID
            if not para_corrections:
                for correction in corrections:
                    original = correction.get('original', '')
                    if original and original in original_para.text:
                        para_corrections.append(correction)
            
            if para_corrections and original_para.text.strip():
                # We have corrections for this paragraph
                original_text = original_para.text
                remaining_text = original_text
                
                # Track if we've applied any corrections
                applied_corrections = False
                
                # Apply each correction
                for correction in para_corrections:
                    if 'original' in correction and 'corrected' in correction:
                        original_part = correction.get('original', '')
                        corrected_part = correction.get('corrected', '')
                        
                        if original_part in remaining_text:
                            # Split the text at the correction point
                            parts = remaining_text.split(original_part, 1)
                            before = parts[0]
                            after = parts[1] if len(parts) > 1 else ""
                            
                            # Add the text before the correction
                            if before:
                                new_para.add_run(before)
                            
                            # Add the correction with visual formatting
                            # Original with strikethrough
                            del_run = new_para.add_run(original_part)
                            del_run.font.strike = True
                            
                            # Corrected in red
                            ins_run = new_para.add_run(corrected_part)
                            ins_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
                            ins_run.bold = True
                                                        
                            # Update remaining text for next correction
                            remaining_text = after
                            applied_corrections = True
                            has_correction = True
                        else:
                            print(f"WARNING: Could not find text '{original_part}' in paragraph text.")
                
                # Add any remaining text
                if remaining_text:
                    new_para.add_run(remaining_text)
                
                # If no corrections were applied, just add the original text
                if not applied_corrections:
                    if len(new_para.runs) == 0:  # Ensure we haven't already added content
                        new_para.text = original_text
            else:
                # No corrections for this paragraph, just copy the text and formatting
                if original_para.runs:
                    # Copy runs to preserve formatting
                    for run in original_para.runs:
                        new_run = new_para.add_run(run.text)
                        # Copy basic formatting
                        try:
                            new_run.bold = run.bold
                            new_run.italic = run.italic
                            new_run.underline = run.underline
                            if hasattr(run, 'font') and hasattr(run.font, 'color') and run.font.color:
                                new_run.font.color.rgb = run.font.color.rgb
                        except Exception as e:
                            print(f"Warning: Could not copy run formatting: {e}")
                else:
                    # If no runs, just copy the text
                    new_para.text = original_para.text
        
        # Copy tables
        copy_tables(original_doc, corrected_doc)
        
        # Save the document
        corrected_doc.save(output_path)
        return output_path
        
    except Exception as e:
        print(f"Error applying corrections to DOCX: {e}")
        import traceback
        traceback.print_exc()
        
        # Create a basic document with corrections
        return create_fallback_document(corrections, output_path)

def copy_docx_formatting(source_doc, target_doc):
    """Copy formatting from source document to target document.
    
    Args:
        source_doc: Source document object
        target_doc: Target document object
    """
    # Copy the base styles
    try:
        for style_id in source_doc.styles:
            if style_id not in target_doc.styles:
                try:
                    style = source_doc.styles[style_id]
                    # We can't directly copy styles between documents,
                    # but we can ensure the style exists in the target document
                    if style.type == 1:  # Paragraph style
                        if style.name not in target_doc.styles:
                            target_doc.styles.add_style(style.name, style.type)
                except Exception as e:
                    print(f"Warning: Could not copy style {style_id}: {e}")
    except Exception as e:
        print(f"Warning: Error copying document styles: {e}")
    
    # Copy the document properties and sections
    for section in source_doc.sections:
        new_section = target_doc.add_section()
        try:
            new_section.page_height = section.page_height
            new_section.page_width = section.page_width
            new_section.left_margin = section.left_margin
            new_section.right_margin = section.right_margin
            new_section.top_margin = section.top_margin
            new_section.bottom_margin = section.bottom_margin
            new_section.header_distance = section.header_distance
            new_section.footer_distance = section.footer_distance
        except Exception as e:
            print(f"Warning: Could not copy section properties: {e}")

def copy_tables(source_doc, target_doc):
    """Copy tables from source document to target document.
    
    Args:
        source_doc: Source document object
        target_doc: Target document object
    """
    for table in source_doc.tables:
        try:
            new_table = target_doc.add_table(rows=len(table.rows), cols=len(table.columns))
            # Copy table style if available
            if hasattr(table, 'style') and table.style:
                try:
                    new_table.style = table.style
                except:
                    print("Warning: Could not copy table style")
            
            # Copy cell contents
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    try:
                        # Copy each paragraph in the cell
                        for para in cell.paragraphs:
                            cell_para = new_table.rows[i].cells[j].add_paragraph()
                            cell_para.text = para.text
                            # Try to copy paragraph style
                            if hasattr(para, 'style') and para.style:
                                try:
                                    cell_para.style = para.style.name
                                except:
                                    pass
                    except Exception as e:
                        # Fall back to simple text copy
                        new_table.rows[i].cells[j].text = cell.text
                        print(f"Warning: Error copying table cell content: {e}")
        except Exception as e:
            print(f"Warning: Could not copy table: {e}")

def create_fallback_document(corrections, output_path):
    """Create a basic document with corrections when main processing fails.
    
    Args:
        corrections: List of correction dictionaries
        output_path: Path to save the document
        
    Returns:
        Path to the saved document
    """
    try:
        from docx import Document
        
        # Convert corrections to list format if needed
        if isinstance(corrections, dict):
            corrections_list = corrections.get('corrections', [])
        elif isinstance(corrections, list):
            corrections_list = corrections
        else:
            corrections_list = [corrections] if corrections else []
        
        # Create a new document
        doc = Document()
        doc.add_heading('Grammar Corrections', 0)
        
        # Add each correction
        for i, corr in enumerate(corrections_list, 1):
            if isinstance(corr, dict):
                p = doc.add_paragraph(f"{i}. Original: ")
                p.add_run(corr.get('original', '')).bold = True
                p.add_run("\nCorrected: ")
                p.add_run(corr.get('corrected', '')).italic = True
                if 'explanation' in corr:
                    p.add_run("\nExplanation: ")
                    p.add_run(corr.get('explanation', ''))
            else:
                doc.add_paragraph(f"{i}. {corr}")
            doc.add_paragraph("")
        
        # Save the document
        doc_path = output_path
        doc.save(doc_path)
        return doc_path
        
    except Exception as e:
        print(f"Error creating fallback document: {e}")
        
        # Create a text file as a last resort
        text_path = f"{output_path}.txt"
        try:
            with open(text_path, 'w', encoding='utf-8') as f:
                f.write("GRAMMAR CORRECTIONS:\n\n")
                for i, corr in enumerate(corrections_list, 1):
                    if isinstance(corr, dict):
                        f.write(f"{i}. Original: {corr.get('original', '')}\n")
                        f.write(f"   Corrected: {corr.get('corrected', '')}\n")
                        if 'explanation' in corr:
                            f.write(f"   Explanation: {corr['explanation']}\n")
                    elif isinstance(corr, str):
                        f.write(f"{i}. Correction: {corr}\n")
                    f.write("\n")
        except Exception as e:
            print(f"Error creating text file: {e}")
            return None
            
        return text_path

def create_structured_docx(content, output_path):
    """Create a structured DOCX document.
    
    Args:
        content: Text content or structured content list
        output_path: Path to save the document
        
    Returns:
        Path to the saved document
    """
    try:
        from docx import Document
        from docx.shared import Pt
        
        doc = Document()
        
        # Process content
        if isinstance(content, str):
            # Simple text content
            paragraphs = content.split('\n\n')
            for para_text in paragraphs:
                if not para_text.strip():
                    continue
                p = doc.add_paragraph(para_text)
        else:
            # Structured content
            for item in content:
                if isinstance(item, dict):
                    item_type = item.get('type', 'paragraph')
                    item_text = item.get('text', '')
                    item_level = item.get('level', 1)
                    
                    if not item_text.strip():
                        continue
                        
                    if item_type == 'heading':
                        doc.add_heading(item_text, level=item_level)
                    else:
                        doc.add_paragraph(item_text)
        
        # Save the document
        doc.save(output_path)
        return output_path
        
    except Exception as e:
        print(f"Error creating structured DOCX: {e}")
        
        # Create a text file as fallback
        text_path = f"{output_path}.txt"
        try:
            with open(text_path, 'w', encoding='utf-8') as f:
                if isinstance(content, str):
                    f.write(content)
                else:
                    for item in content:
                        if isinstance(item, dict):
                            f.write(item.get('text', '') + '\n\n')
                        else:
                            f.write(str(item) + '\n\n')
        except Exception as text_err:
            print(f"Error creating text fallback: {text_err}")
            return None
            
        return text_path