"""Module for reading document files of various formats."""

import os
import re
import zipfile
import io
from lxml import etree
from typing import Tuple, List, Dict, Any, Optional

# Check for docx library availability
DOCX_AVAILABLE = False
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Check for odf library availability and get element classes
ODF_AVAILABLE = False
try:
    from odf.opendocument import OpenDocument
    from odf.text import P, H, Span
    from odf.element import Element as OdfElement
    from odf import teletype
    ODF_AVAILABLE = True
except ImportError:
    ODF_AVAILABLE = False

def read_document(file_path):
    """Read document content and return structured format."""
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.docx':
        content, structured_content = read_docx(file_path)
        return content, structured_content, True
    elif file_ext == '.odt':
        content, structured_content = read_odt(file_path)
        # No need for conversion since read_odt now returns proper format
        return content, structured_content, False
    else:
        content = read_text_file(file_path)
        return content, None, True

def read_docx(file_path):
    """Read content from a DOCX file and return structured representation."""
    try:
        print(f"Reading DOCX file: {file_path}")
        
        doc = Document(file_path)
        full_text = []
        structured_content = []
        element_id = 0
        
        # Extract each paragraph
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                full_text.append(text)
                element_id += 1
                
                # Add to structured content with metadata
                structured_content.append({
                    'id': f'p{element_id}',
                    'type': 'paragraph',
                    'content': text,
                    'style': paragraph.style.name
                })
        
        # Also extract tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text.strip()
                        if text:
                            full_text.append(text)
                            element_id += 1
                            
                            # Add to structured content with metadata
                            structured_content.append({
                                'id': f'p{element_id}',
                                'type': 'table_cell',
                                'content': text,
                                'style': paragraph.style.name
                            })
        
        full_content = '\n\n'.join(full_text)
        print(f"Extracted {len(structured_content)} structured elements from document")
        return full_content, structured_content
        
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
        return "", []

def read_docx_xml(file_path):
    """Read a DOCX document by parsing its XML directly."""
    # ...existing code...

def read_docx_as_zip(file_path):
    """Read a DOCX as a simple ZIP and extract text from document.xml."""
    # ...existing code...

def read_odt(file_path):
    """Read an ODT file and return its content as text."""
    try:
        from odf.opendocument import load
        
        # Load the ODT file
        print(f"Reading ODT file: {file_path}")
        doc = load(file_path)
        
        # Extract all paragraphs and headings
        paragraphs = []
        structured_content = []
        
        # Extract headings and paragraphs
        for i, element in enumerate(doc.text.childNodes):
            # Check if element is an ODT text element
            if isinstance(element, OdfElement) and element.qname[1] in ('p', 'h'):
                content = teletype.extractText(element)
                if content.strip():
                    paragraphs.append(content)
                    structured_content.append({
                        'id': f'p{i+1}',
                        'type': 'heading' if element.qname[1] == 'h' else 'paragraph',
                        'content': content
                    })
        
        # Join paragraphs with newlines
        content = '\n\n'.join(paragraphs)
        
        print(f"Successfully read ODT content ({len(content)} characters)")
        return content, structured_content
        
    except ImportError:
        print("odfpy library not found, using alternative extraction method")
        return read_odt_xml(file_path)

def read_odt_xml(file_path):
    """Read an ODT document by parsing its XML directly."""
    try:
        content = []
        structured_content = []
        
        with zipfile.ZipFile(file_path) as odt:
            with odt.open('content.xml') as f:
                tree = etree.parse(f)
                root = tree.getroot()
                
                # Define namespaces
                ns = {
                    'office': 'urn:oasis:names:tc:opendocument:xmlns:office:1.0',
                    'text': 'urn:oasis:names:tc:opendocument:xmlns:text:1.0'
                }
                
                # Get all paragraphs and headings
                elements = root.xpath('//text:p | //text:h', namespaces=ns)
                
                for i, elem in enumerate(elements):
                    text = elem.text or ''
                    # Add all text from child elements
                    for child in elem.xpath('.//text()', namespaces=ns):
                        text += child
                    
                    if text.strip():
                        content.append(text)
                        elem_type = 'heading' if elem.tag.endswith('h') else 'paragraph'
                        structured_content.append({
                            'id': f'p{i+1}',
                            'type': elem_type,
                            'content': text
                        })
        
        return '\n\n'.join(content), structured_content
        
    except Exception as e:
        print(f"Error reading ODT XML: {e}")
        return read_odt_as_zip(file_path)

def read_odt_as_zip(file_path):
    """Read an ODT as a simple ZIP and extract text from content.xml."""
    # ...existing code...

def read_text_file(file_path):
    """Read a plain text file."""
    # ...existing code...