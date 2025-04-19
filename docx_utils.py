from lxml import etree
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
import uuid
from datetime import datetime

def create_comments_part(document):
    """Create a comments part for the document if it doesn't exist."""
    part = document.part
    
    # Check if comments part already exists
    for rel in part.rels.values():
        if rel.reltype == RT.COMMENTS:
            return rel.target_part
    
    # Create a new comments part
    comments_part_xml = f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"></w:comments>'
    comments_part = part.package.create_part('/word/comments.xml', 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml', comments_part_xml)
    
    # Add relationship to the comments part
    part.rels.add_relationship(RT.COMMENTS, comments_part)
    
    return comments_part

def get_next_comment_id(document):
    """Get the next available comment ID."""
    comments_part = create_comments_part(document)
    comments_element = etree.fromstring(comments_part.blob)
    existing_comments = comments_element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}comment')
    if not existing_comments:
        return 1
    return max([int(comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')) for comment in existing_comments]) + 1

def create_comment_reference(comment_id):
    """Create XML for a comment reference."""
    return f'<w:commentReference w:id="{comment_id}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'

def create_comment(comment_id, text, author, initials=None):
    """Create XML for a comment."""
    from datetime import datetime
    initials = initials or author[0]
    date = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")
    return f'<w:comment w:id="{comment_id}" w:author="{author}" w:initials="{initials}" w:date="{date}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:p><w:r><w:t>{text}</w:t></w:r></w:p></w:comment>'

def copy_run_formatting(source_run, target_run):
    """Copy formatting from one run to another."""
    target_run.bold = source_run.bold
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    if hasattr(source_run, 'font') and hasattr(target_run, 'font'):
        if hasattr(source_run.font, 'size'):
            target_run.font.size = source_run.font.size
        if hasattr(source_run.font, 'name'):
            target_run.font.name = source_run.font.name
        if hasattr(source_run.font, 'color') and hasattr(target_run.font.color, 'rgb'):
            target_run.font.color.rgb = source_run.font.color.rgb

# Adding the add_comment_to_document function from document_writer.py here to avoid circular import
def add_comment_to_document(doc, paragraph, comment_text, author="Grammatik-Prüfung", initials="GP"):
    """Add a comment to a paragraph in a document.
    
    Args:
        doc: The docx Document object
        paragraph: The paragraph to add the comment to
        comment_text: The text of the comment
        author: The author of the comment
        initials: The initials of the author
        
    Returns:
        The comment ID
    """
    # Get or create comments part
    comments_part = ensure_comments_part_exists(doc)
    
    # Create a unique comment ID
    comment_id = str(len(comments_part.xpath('//w:comment')) + 1)
    
    # Create a comment element
    comment = OxmlElement('w:comment')
    comment.set(qn('w:id'), comment_id)
    comment.set(qn('w:author'), author)
    comment.set(qn('w:initials'), initials)
    comment.set(qn('w:date'), datetime.now().isoformat())
    
    # Add the comment text as a paragraph
    comment_para = OxmlElement('w:p')
    comment_r = OxmlElement('w:r')
    comment_text_element = OxmlElement('w:t')
    comment_text_element.text = comment_text
    comment_r.append(comment_text_element)
    comment_para.append(comment_r)
    comment.append(comment_para)
    
    # Add the comment to the comments part
    comments_element = comments_part.getroot()
    comments_element.append(comment)
    
    # Add a comment reference to the paragraph
    run = paragraph.add_run()
    comment_reference = OxmlElement('w:commentReference')
    comment_reference.set(qn('w:id'), comment_id)
    run._element.append(comment_reference)
    
    return comment_id

# Adding the ensure_comments_part_exists function from document_writer.py
def ensure_comments_part_exists(doc):
    """Ensure the document has a comments part.
    
    Args:
        doc: The docx Document object
        
    Returns:
        The comments part
    """
    # Check if document has a main document part
    if not hasattr(doc, 'part'):
        raise ValueError("Document doesn't have a main document part")
    
    # Try to get existing comments part
    try:
        NS = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'comments': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
        }
        for rel in doc.part.rels.values():
            if rel.reltype == NS['comments']:
                return rel.target
    except:
        pass
    
    # Create a new comments part if it doesn't exist
    try:
        from docx.parts.comments import CommentsXml
        comments_content = CommentsXml.new()
        comments_part = doc.part.package.create_part('/word/comments.xml', 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml', comments_content)
        
        # Add a relationship to the comments part
        NS = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
            'comments': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments'
        }
        rel_id = doc.part.relate_to(comments_part, NS['comments'], is_external=False)
        
        # Add comments reference to document settings
        settings = doc.settings
        if settings is None:
            settings = doc.add_settings()
        
        # Return the comments part
        return comments_part
    except Exception as e:
        print(f"Error creating comments part: {e}")
        raise

def add_tracked_change_with_comment(doc, paragraph, original_text, corrected_text, explanation):
    """Add a tracked change and a comment to a paragraph in a DOCX document.

    Args:
        doc: The docx Document object
        paragraph: The paragraph to modify.
        original_text: The original text to replace.
        corrected_text: The corrected text to insert.
        explanation: The explanation for the change, added as a comment.
        
    Returns:
        Boolean indicating success or failure
    """
    try:
        # Setup tracked changes in document settings if not already set
        settings = doc.settings
        if settings is None:
            settings = doc.add_settings()
        
        track_revisions = settings._element.find(qn('w:trackRevisions'))
        if track_revisions is None:
            track_revisions = OxmlElement('w:trackRevisions')
            settings._element.append(track_revisions)
        
        # Get the paragraph's XML
        p_xml = paragraph._element
        
        # Find the run containing the original text
        for run in paragraph.runs:
            if original_text in run.text:
                # Split the run into three parts: before, the match, and after
                before, match, after = run.text.partition(original_text)
                
                # Clear the original run's text
                run.text = before
                
                # Add the <w:del> tag for the original text
                del_run = paragraph.add_run()
                del_element = OxmlElement('w:del')
                del_element.set(qn('w:author'), "Correction")
                del_element.set(qn('w:date'), datetime.now().isoformat())
                del_element.set(qn('w:id'), str(uuid.uuid4())[:8])
                del_run._element.append(del_element)
                
                del_text = OxmlElement('w:t')
                if match.startswith(' ') or match.endswith(' '):
                    del_text.set(qn('xml:space'), 'preserve')
                del_text.text = match
                del_element.append(del_text)
                
                # Add the <w:ins> tag for the corrected text
                ins_run = paragraph.add_run()
                ins_element = OxmlElement('w:ins')
                ins_element.set(qn('w:author'), "Correction")
                ins_element.set(qn('w:date'), datetime.now().isoformat())
                ins_element.set(qn('w:id'), str(uuid.uuid4())[:8])
                ins_run._element.append(ins_element)
                
                ins_text = OxmlElement('w:t')
                if corrected_text.startswith(' ') or corrected_text.endswith(' '):
                    ins_text.set(qn('xml:space'), 'preserve')
                ins_text.text = corrected_text
                ins_element.append(ins_text)
                
                # Add the 'after' part if it exists
                if after:
                    after_run = paragraph.add_run(after)
                    copy_run_formatting(run, after_run)
                
                # Add a comment with the explanation
                add_comment_to_document(doc, paragraph, explanation)
                
                return True
        
        return False
    except Exception as e:
        print(f"Error in add_tracked_change_with_comment: {e}")
        return False